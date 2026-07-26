import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, color_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    # top/bottom: 100 dxa (approx 5pt)
    # left/right: 150 dxa (approx 7.5pt)
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_borders(cell, color="D3D3D3", sz="4", val="single"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), val)
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def prevent_row_split(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

def set_paragraph_shading_and_border(paragraph, color_hex="F8F9FA", border_color_hex="1F4E79"):
    pPr = paragraph._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    pPr.append(shd)
    # Left border of 3pt (sz="24")
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="12" w:color="{border_color_hex}"/></w:pBdr>')
    pPr.append(pBdr)

def add_header_border(header_para):
    pPr = header_para._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="4" w:color="CCCCCC"/></w:pBdr>')
    pPr.append(pBdr)

def add_page_number(run):
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'PAGE')
    run._r.append(fldSimple)

def parse_and_add_text(paragraph, text, default_font_size=11, is_code=False):
    # Regex to extract formatting: ***bold-italic***, **bold**, *italic*, `code`, [text](url)
    pattern = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))')
    parts = pattern.split(text)
    
    for part in parts:
        if not part:
            continue
        
        if part.startswith('***') and part.endswith('***'):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(default_font_size - 1)
        elif part.startswith('[') and ']' in part and '(' in part and part.endswith(')'):
            match = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if match:
                link_text, url = match.groups()
                run = paragraph.add_run(link_text)
                run.font.color.rgb = RGBColor(41, 128, 185)  # Soft blue
                run.underline = True
            else:
                run = paragraph.add_run(part)
        else:
            run = paragraph.add_run(part)
            
        if not part.startswith('`'):
            run.font.name = 'Times New Roman'
            run.font.size = Pt(default_font_size)
            if is_code:
                run.font.name = 'Courier New'
                run.font.size = Pt(default_font_size - 1.5)

def join_paragraph_lines(lines):
    result = []
    for line in lines:
        if line.endswith('  '):
            result.append(line.rstrip() + '\n')
        else:
            result.append(line)
    
    text = ""
    for i, line in enumerate(result):
        if i == 0:
            text = line
        else:
            if text.endswith('\n'):
                text += line
            else:
                text += " " + line
    return text

def get_column_widths(num_cols):
    if num_cols == 2:
        return [Inches(2.0), Inches(4.5)]
    elif num_cols == 3:
        return [Inches(1.8), Inches(2.7), Inches(2.0)]
    elif num_cols == 9:
        return [Inches(1.3)] + [Inches(0.65)] * 8
    else:
        return [Inches(6.5 / num_cols)] * num_cols

def parse_markdown(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    
    if content.startswith('\ufeff'):
        content = content[1:]
        
    lines = content.splitlines()
    
    blocks = []
    current_block = None
    in_code_block = False
    code_content = []
    
    for line in lines:
        stripped = line.strip()
        
        # 1. Code block handling
        if stripped.startswith('```'):
            if in_code_block:
                blocks.append({
                    'type': 'code_block',
                    'content': '\n'.join(code_content)
                })
                in_code_block = False
                code_content = []
            else:
                if current_block:
                    blocks.append(current_block)
                    current_block = None
                in_code_block = True
            continue
            
        if in_code_block:
            code_content.append(line)
            continue
            
        # 2. Empty line handling
        if not stripped:
            if current_block:
                blocks.append(current_block)
                current_block = None
            continue
            
        # 3. Headings
        if stripped.startswith('# '):
            if current_block: blocks.append(current_block)
            blocks.append({'type': 'title', 'content': stripped[2:]})
            current_block = None
            continue
        elif stripped.startswith('## '):
            if current_block: blocks.append(current_block)
            blocks.append({'type': 'h1', 'content': stripped[3:]})
            current_block = None
            continue
        elif stripped.startswith('### '):
            if current_block: blocks.append(current_block)
            blocks.append({'type': 'h2', 'content': stripped[4:]})
            current_block = None
            continue
        elif stripped.startswith('#### '):
            if current_block: blocks.append(current_block)
            blocks.append({'type': 'h3', 'content': stripped[5:]})
            current_block = None
            continue
            
        # 4. Blockquotes
        if stripped.startswith('> '):
            if current_block and current_block['type'] != 'blockquote':
                blocks.append(current_block)
                current_block = None
            
            val = stripped[2:]
            if current_block:
                current_block['content'].append(val)
            else:
                current_block = {'type': 'blockquote', 'content': [val]}
            continue
            
        # 5. Tables
        if stripped.startswith('|'):
            if current_block and current_block['type'] != 'table':
                blocks.append(current_block)
                current_block = None
            
            row_cells = [c.strip() for c in stripped.split('|')[1:-1]]
            
            # Check if separator row
            is_separator = False
            if row_cells and all(re.match(r'^:?-+:?$', c) for c in row_cells):
                is_separator = True
                
            if not is_separator:
                if current_block:
                    current_block['content'].append(row_cells)
                else:
                    current_block = {'type': 'table', 'content': [row_cells]}
            continue
            
        # 6. Bullet lists
        if stripped.startswith('- ') or stripped.startswith('* '):
            if current_block and current_block['type'] != 'bullet_list':
                blocks.append(current_block)
                current_block = None
                
            val = stripped[2:]
            if current_block:
                current_block['content'].append(val)
            else:
                current_block = {'type': 'bullet_list', 'content': [val]}
            continue
            
        # 7. Numbered lists
        match_num = re.match(r'^(\d+)\.\s(.*)', stripped)
        if match_num:
            if current_block and current_block['type'] != 'numbered_list':
                blocks.append(current_block)
                current_block = None
                
            val = match_num.group(2)
            if current_block:
                current_block['content'].append(val)
            else:
                current_block = {'type': 'numbered_list', 'content': [val]}
            continue
            
        # 8. Paragraphs
        if current_block and current_block['type'] != 'paragraph':
            blocks.append(current_block)
            current_block = None
            
        if current_block:
            current_block['content'].append(line)
        else:
            current_block = {'type': 'paragraph', 'content': [line]}
            
    if current_block:
        blocks.append(current_block)
        
    return blocks

def generate_report():
    doc = Document()
    
    # 1. Page Margins Setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Configure Header & Footer
        # Header text
        header = section.header
        header_p = header.paragraphs[0]
        header_p.text = ""
        run = header_p.add_run("Hệ thống University SMS — Tài liệu Nghiệp vụ & Hướng dẫn Sử dụng")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(8.5)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_header_border(header_p)
        
        # Footer text
        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.text = ""
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_p.add_run("Trang ")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
        add_page_number(run)
        
    # Configure styles
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)  # Charcoal/navy tone
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)
    
    # 2. Cover Page
    for _ in range(3):
        doc.add_paragraph()
        
    # Institution
    p_inst = doc.add_paragraph()
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_inst = p_inst.add_run("PHÂN HỆ QUẢN LÝ THÔNG TIN SINH VIÊN (UNIVERSITY SMS)\nHỆ THỐNG PHẦN MỀM ODOO v17")
    run_inst.font.name = 'Times New Roman'
    run_inst.font.size = Pt(12)
    run_inst.bold = True
    run_inst.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    p_inst.paragraph_format.space_after = Pt(100)
    
    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("BÁO CÁO NGHIỆP VỤ CHI TIẾT\nVÀ HƯỚNG DẪN SỬ DỤNG HỆ THỐNG")
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(24)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)  # Deep Navy Blue
    p_title.paragraph_format.space_after = Pt(20)
    
    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run(
        "Tài liệu nghiệp vụ phân tích chi tiết quy trình, vai trò phân quyền hệ thống,\n"
        "dữ liệu mẫu và hướng dẫn thao tác vận hành hoàn chỉnh dành cho người dùng."
    )
    run_sub.font.name = 'Times New Roman'
    run_sub.font.size = Pt(11)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)
    p_sub.paragraph_format.space_after = Pt(120)
    
    # Metadata block using custom tab stops or a simple centered block
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = p_meta.add_run(
        "Phiên bản: 1.0 (Bản chính thức)\n"
        "Ngày cập nhật: 08 tháng 07 năm 2026\n"
        "Phạm vi: Các module odoo_sms_*\n"
        "Ban biên soạn: Đội ngũ phát triển University SMS"
    )
    run_meta.font.name = 'Times New Roman'
    run_meta.font.size = Pt(10.5)
    run_meta.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    
    p_loc = doc.add_paragraph()
    p_loc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_loc.paragraph_format.space_before = Pt(80)
    run_loc = p_loc.add_run("Hà Nội, Năm 2026")
    run_loc.font.name = 'Times New Roman'
    run_loc.font.size = Pt(11)
    run_loc.bold = True
    run_loc.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    
    doc.add_page_break()
    
    # 3. Table of Contents Placeholder
    p_toc_head = doc.add_paragraph()
    p_toc_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_toc = p_toc_head.add_run("MỤC LỤC TÀI LIỆU")
    run_toc.font.name = 'Times New Roman'
    run_toc.font.size = Pt(14)
    run_toc.bold = True
    run_toc.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    p_toc_head.paragraph_format.space_after = Pt(12)
    
    p_toc_note = doc.add_paragraph()
    run_note = p_toc_note.add_run(
        "[Hướng dẫn: Để cập nhật Mục lục trong Microsoft Word, vui lòng nhấn chuột phải vào vùng mục lục bên dưới và chọn 'Update Field']"
    )
    run_note.font.italic = True
    run_note.font.size = Pt(9.5)
    run_note.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    p_toc_note.paragraph_format.space_after = Pt(12)
    
    # Insert TOC XML
    p_toc = doc.add_paragraph()
    run_toc_field = p_toc.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run_toc_field._r.append(fldChar1)
    run_toc_field._r.append(instrText)
    run_toc_field._r.append(fldChar2)
    run_toc_field._r.append(fldChar3)
    
    doc.add_page_break()
    
    # 4. Parse & Render Content
    md_file = os.path.join("docs", "TAI_LIEU_NGHIEP_VU_USER_GUIDE_UNIV_SMS.md")
    if not os.path.exists(md_file):
        md_file = "TAI_LIEU_NGHIEP_VU_USER_GUIDE_UNIV_SMS.md"
        
    blocks = parse_markdown(md_file)
    
    for block in blocks:
        b_type = block['type']
        
        # Skip top title as we created a cover page
        if b_type == 'title':
            continue
            
        elif b_type == 'h1':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(block['content'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(15)
            run.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79) # Primary Deep Blue
            
        elif b_type == 'h2':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(block['content'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12.5)
            run.bold = True
            run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6) # Secondary Blue
            
        elif b_type == 'h3':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(block['content'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.bold = True
            run.font.color.rgb = RGBColor(0x40, 0x40, 0x40) # Dark Gray
            
        elif b_type == 'paragraph':
            text = join_paragraph_lines(block['content'])
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Check if this paragraph is a special note (starts with "Lưu ý:" or "Ghi chú:")
            stripped_text = text.strip()
            if stripped_text.startswith("Lưu ý:") or stripped_text.startswith("Ghi chú:"):
                # Style as a callout box
                set_paragraph_shading_and_border(p, color_hex="FDF2E9", border_color_hex="D35400") # Warm orange accent
                parse_and_add_text(p, text, default_font_size=10.5)
            else:
                parse_and_add_text(p, text, default_font_size=11)
                
        elif b_type == 'bullet_list':
            for item in block['content']:
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                parse_and_add_text(p, item, default_font_size=11)
                
        elif b_type == 'numbered_list':
            for item in block['content']:
                p = doc.add_paragraph(style='List Number')
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                parse_and_add_text(p, item, default_font_size=11)
                
        elif b_type == 'blockquote':
            joined_text = "\n".join(block['content'])
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.line_spacing = 1.15
            set_paragraph_shading_and_border(p, color_hex="F2F4F4", border_color_hex="7F8C8D") # Neutral gray callout
            parse_and_add_text(p, joined_text, default_font_size=10.5)
            
        elif b_type == 'code_block':
            # Create paragraph for the entire block
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.line_spacing = 1.0
            set_paragraph_shading_and_border(p, color_hex="F8F9FA", border_color_hex="2C3E50") # Dark slate accent
            parse_and_add_text(p, block['content'], default_font_size=9.5, is_code=True)
            
        elif b_type == 'table':
            rows_data = block['content']
            if not rows_data:
                continue
                
            num_cols = len(rows_data[0])
            num_rows = len(rows_data)
            
            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            widths = get_column_widths(num_cols)
            
            # Apply styling & set content
            for row_idx, row_data in enumerate(rows_data):
                row = table.rows[row_idx]
                prevent_row_split(row)
                
                is_header = (row_idx == 0)
                if is_header:
                    repeat_header(row)
                    
                # We need to handle column data, some rows might have fewer cells than header
                # so check safety:
                for col_idx, cell_value in enumerate(row_data):
                    if col_idx >= num_cols:
                        break
                    cell = row.cells[col_idx]
                    
                    # Set cell content
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.line_spacing = 1.1
                    
                    if is_header:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run(cell_value)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)  # White text
                        set_cell_background(cell, "1F4E79")  # Primary Deep Blue
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        # If value looks like a link or status, we can center it
                        if cell_value.startswith('http') or cell_value.startswith('`http'):
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif cell_value in ['CRUD', 'Không', 'Có mặt', 'Vắng mặt', 'Có phép', 'Đi trễ']:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                        parse_and_add_text(p, cell_value, default_font_size=9.5)
                        
                        # Alternating row background shading
                        if row_idx % 2 == 1:
                            set_cell_background(cell, "F2F4F8")  # Light gray-blue
                        else:
                            set_cell_background(cell, "FFFFFF")
                            
                    # Set cell padding and border
                    set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
                    set_cell_borders(cell, color="D3D3D3", sz="4", val="single")
                    
                    # Set column width
                    cell.width = widths[col_idx]
                    
            # Space after table
            p_spacer = doc.add_paragraph()
            p_spacer.paragraph_format.space_before = Pt(0)
            p_spacer.paragraph_format.space_after = Pt(6)
            
    # Save the document
    out_dir = "docs"
    if not os.path.exists(out_dir):
        os.makedirs(out_dir)
    out_path = os.path.join(out_dir, "Bao_Cao_Nghiep_Vu_University_SMS.docx")
    doc.save(out_path)
    print("DOCX report generated successfully at: " + out_path)

if __name__ == "__main__":
    generate_report()