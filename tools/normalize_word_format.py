from __future__ import annotations

import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
INPUT = (
    Path(sys.argv[1]).resolve()
    if len(sys.argv) > 1
    else ROOT / "docs" / "BAO_CAO_DO_AN_TOT_NGHIEP_UNIV_SMS_DINH_DANG_CHUAN_normalized.docx"
)
OUTPUT = (
    Path(sys.argv[2]).resolve()
    if len(sys.argv) > 2
    else ROOT / "docs" / "BAO_CAO_DO_AN_TOT_NGHIEP_UNIV_SMS_DONG_BO_FONT.docx"
)

BODY_SIZE = 13
FONT_NAME = "Times New Roman"


def set_rfonts(rpr, font_name=FONT_NAME):
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = rpr._add_rFonts()
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def set_run_font(run, size=BODY_SIZE):
    run.font.name = FONT_NAME
    run._element.get_or_add_rPr()
    set_rfonts(run._element.rPr)
    run.font.size = Pt(size)


def set_style_font(style, size=BODY_SIZE, bold=None, italic=None):
    style.font.name = FONT_NAME
    style._element.get_or_add_rPr()
    set_rfonts(style._element.rPr)
    style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic


def set_paragraph_spacing(paragraph, line_spacing=1.5, before=0, after=0, justify=True):
    fmt = paragraph.paragraph_format
    fmt.line_spacing = line_spacing
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    if justify:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def configure_styles(doc):
    for style in doc.styles:
        if getattr(style, "type", None) is None:
            continue
        if not hasattr(style, "font"):
            continue
        name = style.name
        if name == "Heading 1":
            set_style_font(style, 14, bold=True, italic=False)
            style.paragraph_format.line_spacing = 1.5
            style.paragraph_format.space_before = Pt(12)
            style.paragraph_format.space_after = Pt(6)
        elif name in {"Heading 2", "Heading 3"}:
            set_style_font(style, BODY_SIZE, bold=True, italic=(name == "Heading 2"))
            style.paragraph_format.line_spacing = 1.5
            style.paragraph_format.space_before = Pt(10)
            style.paragraph_format.space_after = Pt(4)
        elif name in {"Caption", "CaptionFigure", "CaptionTable"}:
            set_style_font(style, BODY_SIZE, italic=True)
            style.paragraph_format.line_spacing = 1.0
            style.paragraph_format.space_before = Pt(3)
            style.paragraph_format.space_after = Pt(3)
        elif "Code" in name:
            set_style_font(style, BODY_SIZE)
            style.paragraph_format.line_spacing = 1.0
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(0)
        else:
            set_style_font(style, BODY_SIZE)
            if hasattr(style, "paragraph_format"):
                style.paragraph_format.line_spacing = 1.5
                style.paragraph_format.space_before = Pt(0)
                style.paragraph_format.space_after = Pt(0)


def paragraph_has_non_text_content(paragraph):
    xml = paragraph._p.xml
    markers = (
        "<w:br",
        "<w:drawing",
        "<w:pict",
        "<w:object",
        "<w:fldChar",
        "<w:instrText",
        "<w:sectPr",
        "<w:lastRenderedPageBreak",
    )
    return any(marker in xml for marker in markers)


def is_blank_paragraph(paragraph):
    return not paragraph.text.strip() and not paragraph_has_non_text_content(paragraph)


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def collapse_blank_paragraphs(doc, max_consecutive=1):
    removed = 0
    blank_run = 0
    for paragraph in list(doc.paragraphs):
        if is_blank_paragraph(paragraph):
            blank_run += 1
            if blank_run > max_consecutive:
                remove_paragraph(paragraph)
                removed += 1
        else:
            blank_run = 0
    return removed


def all_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
    for section in doc.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                yield paragraph


def apply_paragraph_formatting(doc):
    for paragraph in all_paragraphs(doc):
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if style_name == "Heading 1":
            set_paragraph_spacing(paragraph, 1.5, before=12, after=6, justify=False)
        elif style_name in {"Heading 2", "Heading 3"}:
            set_paragraph_spacing(paragraph, 1.5, before=10, after=4, justify=False)
        elif "Caption" in style_name:
            set_paragraph_spacing(paragraph, 1.0, before=3, after=3, justify=False)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "Code" in style_name:
            set_paragraph_spacing(paragraph, 1.0, before=0, after=0, justify=False)
        elif style_name.startswith("toc "):
            set_paragraph_spacing(paragraph, 1.15, before=0, after=0, justify=False)
        else:
            set_paragraph_spacing(paragraph, 1.5, before=0, after=0, justify=True)

        for run in paragraph.runs:
            size = 14 if style_name == "Heading 1" else BODY_SIZE
            set_run_font(run, size=size)


def configure_sections(doc):
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)


def normalize():
    if not INPUT.exists():
        raise FileNotFoundError(INPUT)
    doc = Document(INPUT)
    configure_sections(doc)
    configure_styles(doc)
    removed = collapse_blank_paragraphs(doc, max_consecutive=0)
    apply_paragraph_formatting(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)
    print(f"Removed blank paragraphs: {removed}")


if __name__ == "__main__":
    normalize()
