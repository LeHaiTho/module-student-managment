from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs" / "BAO_CAO_DO_AN_TOT_NGHIEP_UNIV_SMS.md"
OUT_PATH = Path(sys.argv[1]).resolve() if len(sys.argv) > 1 else ROOT / "docs" / "BAO_CAO_DO_AN_TOT_NGHIEP_UNIV_SMS.docx"
SCREENSHOT_DIR = ROOT / "docs"
DIAGRAM_DIR = ROOT / "docs" / "diagrams"


SPECIAL_HEADINGS = {"MỤC LỤC", "DANH MỤC HÌNH", "DANH MỤC BẢNG"}
BODY_FONT_SIZE = 13


def set_font(run, name="Times New Roman", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def shade_element(element, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    element.get_or_add_tcPr().append(shd)


def add_field(paragraph, instruction, placeholder=""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    if placeholder:
        out = paragraph.add_run(placeholder)
        set_font(out)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def add_page_number_field(paragraph):
    add_field(paragraph, "PAGE")


def configure_section_geometry(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)


def start_page_numbering(section, start=1):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))


def add_center_footer_page_number(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    for paragraph in section.header.paragraphs:
        paragraph.clear()
    for paragraph in section.footer.paragraphs:
        paragraph.clear()
    paragraph = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number_field(paragraph)
    for run in paragraph.runs:
        set_font(run, size=BODY_FONT_SIZE)


def add_tc_field(paragraph, text, flag):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    safe_text = text.replace('"', "'")
    instr.text = f'TC "{safe_text}" \\f {flag}'
    run._r.append(instr)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def force_update_fields(doc):
    settings = doc.settings.element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


def clean_inline(text):
    text = text.strip()
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = text.replace("**", "").replace("__", "")
    text = text.replace("`", "")
    text = text.replace("<br>", "\n").replace("<br/>", "\n")
    return text


def extract_figure_caption(line):
    stripped = line.strip()
    match = re.match(r"^\*\*(Hình\s+\d+(?:\.\d+)?\.\s+.*?)\*\*\s*$", stripped)
    if match:
        return clean_inline(match.group(1))
    return ""


def caption_diagram_path(caption, index=0):
    match = re.search(r"Hình\s+(\d+)\.(\d+)", caption, flags=re.IGNORECASE)
    if match:
        return DIAGRAM_DIR / f"hinh-{match.group(1)}-{match.group(2)}-mermaid.png"
    return DIAGRAM_DIR / f"mermaid-{index:02d}.png"


def find_next_figure_caption(lines, close_index):
    for line in lines[close_index + 1 : close_index + 9]:
        caption = extract_figure_caption(line)
        if caption:
            return caption
    return ""


def normalize_heading_text(text):
    text = clean_inline(text)
    return re.sub(r"^\d+(\.\d+)*\s+", "", text).strip()


def is_table_separator(line):
    stripped = line.strip()
    if not stripped.startswith("|"):
        return False
    cells = [c.strip() for c in stripped.strip("|").split("|")]
    return all(re.fullmatch(r":?-{3,}:?", c or "") for c in cells)


def parse_table(lines):
    rows = []
    for line in lines:
        if is_table_separator(line):
            continue
        cells = [clean_inline(c) for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    width = max((len(r) for r in rows), default=0)
    for row in rows:
        row.extend([""] * (width - len(row)))
    return rows


def add_caption(doc, text, kind):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if kind == "figure" else WD_ALIGN_PARAGRAPH.LEFT
    add_tc_field(p, text, "F" if kind == "figure" else "T")
    run = p.add_run(text)
    set_font(run, size=11, italic=True)
    return p


def add_picture_with_caption(doc, image_path, caption, missing_label=None):
    doc.add_paragraph()
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        try:
            run.add_picture(str(image_path), width=Inches(6.25))
        except Exception:
            note = doc.add_paragraph()
            label = missing_label or str(image_path)
            r = note.add_run(f"[Không chèn được ảnh: {label}]")
            set_font(r, italic=True, color=(180, 0, 0))
    else:
        note = doc.add_paragraph()
        label = missing_label or str(image_path)
        r = note.add_run(f"[Thiếu file ảnh: {label}]")
        set_font(r, italic=True, color=(180, 0, 0))
    if caption:
        add_caption(doc, caption, "figure")


def add_mermaid_figure(doc, code_lines, caption, index):
    image_path = caption_diagram_path(caption, index)
    rel = image_path.relative_to(ROOT) if image_path.exists() else image_path
    add_picture_with_caption(doc, image_path, caption or f"Hình Mermaid {index}", str(rel))


def add_table(doc, rows, caption=None):
    if not rows:
        return
    if caption:
        add_caption(doc, caption, "table")
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = True
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(cell_text)
            set_font(run, size=10 if len(rows[0]) > 3 else 11, bold=(r_idx == 0))
            if r_idx == 0:
                shade_element(cell._tc, "D9EAF7")
    doc.add_paragraph()


def add_code_block(doc, code_lines, lang=""):
    if lang:
        p = doc.add_paragraph(style="CodeBlock")
        r = p.add_run(f"# {lang}")
        set_font(r, name="Consolas", size=8.5, color=(90, 90, 90))
    for line in code_lines:
        p = doc.add_paragraph(style="CodeBlock")
        r = p.add_run(line.rstrip() or " ")
        set_font(r, name="Consolas", size=8.5)


def parse_screenshot_rows(lines):
    rows = parse_table(lines)
    out = []
    for row in rows[1:]:
        if len(row) < 7:
            continue
        if not row[0].startswith("Hình "):
            continue
        out.append(
            {
                "figure": row[0],
                "screen": row[1],
                "file": row[2],
                "explain": row[3],
                "meaning": row[4],
                "role": row[5],
                "flow": row[6],
            }
        )
    return out


def add_screenshot_entry(doc, item):
    rel = item["file"].replace("\\", "/")
    img_path = SCREENSHOT_DIR / rel
    add_picture_with_caption(doc, img_path, f"{item['figure']}. {item['screen']}", rel)
    details = (
        f"Giải thích: {item['explain']}. "
        f"Ý nghĩa: {item['meaning']}. "
        f"Vai trò: {item['role']}. "
        f"Luồng xử lý: {item['flow']}."
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(details)
    set_font(r, size=BODY_FONT_SIZE)


def add_paragraph_from_markdown(doc, line):
    stripped = line.strip()
    if not stripped:
        doc.add_paragraph()
        return
    m_img = re.match(r"^!\[([^\]]+)\]\(([^)]+)\)\s*$", stripped)
    if m_img:
        caption = clean_inline(m_img.group(1))
        rel = m_img.group(2).replace("\\", "/").strip()
        img_path = SCREENSHOT_DIR / rel
        if img_path.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(img_path), width=Inches(6.25))
        else:
            note = doc.add_paragraph()
            r = note.add_run(f"[Thiếu file ảnh: {rel}]")
            set_font(r, italic=True, color=(180, 0, 0))
        add_caption(doc, caption, "figure")
        return
    if stripped.startswith(">"):
        p = doc.add_paragraph(style="Intense Quote")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(clean_inline(stripped[1:].strip()))
        set_font(r, size=BODY_FONT_SIZE, italic=True)
        return
    if re.match(r"^- ", stripped):
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(clean_inline(stripped[2:]))
        set_font(r, size=BODY_FONT_SIZE)
        return
    m_num = re.match(r"^(\d+)\.\s+(.*)$", stripped)
    if m_num:
        p = doc.add_paragraph(style="List Number")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(clean_inline(m_num.group(2)))
        set_font(r, size=BODY_FONT_SIZE)
        return
    caption = extract_figure_caption(stripped)
    if caption:
        add_caption(doc, caption, "figure")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(clean_inline(stripped))
    set_font(r, size=BODY_FONT_SIZE)


def add_cover(doc, lines):
    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        if stripped.startswith("#"):
            text = clean_inline(stripped.lstrip("#").strip())
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            size = 16 if "BÁO CÁO" in text or "ĐỀ TÀI" in text else 14
            set_font(r, size=size, bold=True)
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(clean_inline(stripped))
            set_font(r, size=BODY_FONT_SIZE)
    doc.add_page_break()


def configure_document(doc):
    section = doc.sections[0]
    configure_section_geometry(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(BODY_FONT_SIZE)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    heading_specs = {
        "Heading 1": (14, False),
        "Heading 2": (13, True),
        "Heading 3": (13, False),
    }
    for style_name, (size, italic) in heading_specs.items():
        st = styles[style_name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.italic = italic
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.line_spacing = 1.5

    caption = styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    caption.font.size = Pt(11)
    caption.font.italic = True

    for name in ["CaptionFigure", "CaptionTable"]:
        if name not in styles:
            st = styles.add_style(name, 1)
        else:
            st = styles[name]
        st.base_style = styles["Caption"]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.font.size = Pt(11)
        st.font.italic = True

    if "CodeBlock" not in styles:
        code = styles.add_style("CodeBlock", 1)
    else:
        code = styles["CodeBlock"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code.font.size = Pt(9)
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after = Pt(0)


def build():
    text = MD_PATH.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    configure_document(doc)

    first_break = lines.index("\\pagebreak") if "\\pagebreak" in lines else 0
    add_cover(doc, lines[:first_break])

    i = first_break + 1
    skip_special = False
    in_code = False
    code_lang = ""
    code_lines = []
    table_lines = []
    current_chapter = 0
    table_counters = {}
    nearest_heading = ""
    mermaid_counter = 0
    skip_figure_captions = set()
    last_heading_upper = ""

    def flush_table():
        nonlocal table_lines
        if not table_lines:
            return
        if table_lines and "screenshots/" in "\n".join(table_lines):
            for item in parse_screenshot_rows(table_lines):
                add_screenshot_entry(doc, item)
        else:
            rows = parse_table(table_lines)
            caption = None
            if current_chapter > 0:
                table_counters[current_chapter] = table_counters.get(current_chapter, 0) + 1
                caption = f"Bảng {current_chapter}.{table_counters[current_chapter]}. {nearest_heading}"
            add_table(doc, rows, caption=caption)
        table_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if in_code:
            if stripped.startswith("```"):
                if code_lang.lower() == "mermaid":
                    mermaid_counter += 1
                    caption = find_next_figure_caption(lines, i)
                    add_mermaid_figure(doc, code_lines, caption, mermaid_counter)
                    if caption:
                        skip_figure_captions.add(caption)
                else:
                    add_code_block(doc, code_lines, code_lang)
                in_code = False
                code_lines = []
                code_lang = ""
            else:
                code_lines.append(line)
            i += 1
            continue

        if table_lines and not stripped.startswith("|"):
            flush_table()

        if stripped == "\\pagebreak":
            flush_table()
            if last_heading_upper == "MỤC LỤC":
                section = doc.add_section(WD_SECTION.NEW_PAGE)
                configure_section_geometry(section)
                start_page_numbering(section, 1)
                add_center_footer_page_number(section)
            else:
                doc.add_page_break()
            skip_special = False
            i += 1
            continue

        if stripped.startswith("```"):
            flush_table()
            in_code = True
            code_lang = stripped.strip("`").strip()
            code_lines = []
            i += 1
            continue

        heading = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if heading:
            flush_table()
            level = len(heading.group(1))
            text = clean_inline(heading.group(2))
            last_heading_upper = text.upper()
            nearest_heading = normalize_heading_text(text)
            chap = re.match(r"CHƯƠNG\s+(\d+)", text, flags=re.IGNORECASE)
            if chap:
                current_chapter = int(chap.group(1))
            elif text.startswith("PHỤ LỤC") or text in {"LỜI CẢM ƠN", "NHẬN XÉT GIẢNG VIÊN", "LỜI MỞ ĐẦU", "KẾT LUẬN"}:
                current_chapter = 0
            p = doc.add_paragraph(style=f"Heading {level}")
            r = p.add_run(text)
            set_font(r, size={1: 14, 2: 13, 3: 13}[level], bold=True, italic=(level == 2))

            upper = text.upper()
            if upper in SPECIAL_HEADINGS:
                if upper == "MỤC LỤC":
                    p2 = doc.add_paragraph()
                    add_field(p2, r'TOC \o "1-3" \h \z \u', "Bấm Ctrl+A rồi F9 trong Word để cập nhật mục lục.")
                elif upper == "DANH MỤC HÌNH":
                    p2 = doc.add_paragraph()
                    add_field(p2, r'TOC \h \z \f F', "Bấm Ctrl+A rồi F9 trong Word để cập nhật danh mục hình.")
                elif upper == "DANH MỤC BẢNG":
                    p2 = doc.add_paragraph()
                    add_field(p2, r'TOC \h \z \f T', "Bấm Ctrl+A rồi F9 trong Word để cập nhật danh mục bảng.")
                skip_special = True
            else:
                skip_special = False
            i += 1
            continue

        if skip_special:
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines.append(line)
            i += 1
            continue

        caption = extract_figure_caption(stripped)
        if caption and caption in skip_figure_captions:
            skip_figure_captions.remove(caption)
            i += 1
            continue

        add_paragraph_from_markdown(doc, line)
        i += 1

    flush_table()
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
